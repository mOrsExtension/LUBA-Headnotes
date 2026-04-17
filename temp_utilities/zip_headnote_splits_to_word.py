import copy
from pathlib import Path
from docx import Document
from natsort import natsorted

# zips together docx files in folder into combined file(s) (kept alphabetically by file name)

# ── CONFIG ──────────────────────────────────────────────────────────────────
FOLDER_PATH = r"..\..\..\Word Versions"
OUTPUT_NAME = "combined.docx"              # output file, saved folder
# ────────────────────────────────────────────────────────────────────────────

def combine_documents(folder_path):
    folder_path = Path(folder_path)
    files = natsorted(folder_path.glob("*.docx"), key=lambda f: f.name)

    # Exclude the output file itself in case it already exists in the folder
    files = [f for f in files if f.name != OUTPUT_NAME]

    if not files:
        print("No .docx files found.")
        return

    print(f"Combining {len(files)} files...")

    combined = Document()
    # Remove the default empty paragraph Word adds to new documents
    for p in combined.paragraphs:
        p._element.getparent().remove(p._element)

    for path in files:
        doc = Document(path)
        for para in doc.paragraphs:
            combined.element.body.append(copy.deepcopy(para._element))
        print(f"  Added: {path.name}")

    out_path = folder_path / OUTPUT_NAME
    combined.save(out_path)
    print(f"\nDone. Saved to: {out_path}")

combine_documents(FOLDER_PATH)