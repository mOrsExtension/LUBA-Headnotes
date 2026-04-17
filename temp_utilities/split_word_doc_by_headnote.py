import re
import copy
from pathlib import Path
from docx import Document
# Takes Word document full of headnotes and splits them into smaller files for each headnote (1.1.1; 1.1.2, etc.), provided that the first paragraph starts with headnote number in bold face.

# ── CONFIG ──────────────────────────────────────────────────────────────────
INPUT_PATH = r"..\headnotes_full-2025-08-06.docx"
# ────────────────────────────────────────────────────────────────────────────

# Matches headnotes like: 1.1.1  16.  3.2  10.1.2.4
HEADNOTE_RE = re.compile(r"^(\d{1,2}\.)+(\d{1,2})?$")


def extract_headnote(para):
    """
    Return the headnote label (e.g. '1.1.1') if this paragraph starts with one, else None.
    Checks the start of the paragraph's full text against HEADNOTE_RE, ignoring formatting.
    """
    text = para.text.strip()
    if not text:
        return None
    token = text.split()[0].rstrip(".")
    if HEADNOTE_RE.match(token + "."):
        return token
    return None


def copy_paragraph_to_doc(src_para, dest_doc):
    """Deep-copy a paragraph's XML into the destination document body."""
    new_para = copy.deepcopy(src_para._element)
    dest_doc.element.body.append(new_para)

def safe_filename(label):
    """E.g. turn  '1.1.1' into a '1_1_1' for safe filename component"""
    return label.replace(".", "_")

def split_document(input_path):
    input_path = Path(input_path)
    doc = Document(input_path)
    paragraphs = doc.paragraphs

    # Walk paragraphs, recording where each *new* headnote label begins.
    # Multiple consecutive paragraphs with the same label are grouped together.
    splits = []   # list of (start_index, label)
    current_label = None

    for i, para in enumerate(paragraphs):
        label = extract_headnote(para)
        if label and label != current_label:
            splits.append((i, label))
            current_label = label

    if not splits:
        print("No headnotes found. Check that the bold prefix text matches the expected pattern.")
        return

    print(f"Found {len(splits)} unique headnote sections.")

    # Build ranges: each section runs from its start up to the next split point
    ranges = []
    for idx, (start, label) in enumerate(splits):
        end = splits[idx + 1][0] if idx + 1 < len(splits) else len(paragraphs)
        ranges.append((start, end, label))

    output_dir = input_path.parent

    for i, (start, end, label) in enumerate(ranges):
        new_doc = Document()
        # Remove the default empty paragraph Word adds to new documents
        for p in new_doc.paragraphs:
            p._element.getparent().remove(p._element)

        for para in paragraphs[start:end]:
            copy_paragraph_to_doc(para, new_doc)

        filename = f"{input_path.stem}_{str(i+1).zfill(3)}_{safe_filename(label)}.docx"
        out_path = output_dir / filename
        new_doc.save(out_path)
        print(f"  Saved: {filename}  ({end - start} paragraphs)")

    print("Done.")

split_document(INPUT_PATH)